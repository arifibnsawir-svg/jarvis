"""
DOCX renderer — bangun naskah panjang dari spec JSON (PIPA2 drafting).
Deterministik. Return artifact_facts utk GUARDIAN.

Dep: pip install python-docx
"""
from __future__ import annotations

import os
from spec_schema import validate_spec, extract_text

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def render_docx(spec: dict, out_path: str) -> dict:
    spec = validate_spec(spec)
    doc = Document()

    # style dasar
    normal = doc.styles["Normal"]
    normal.font.name = spec.get("font", "Calibri")
    normal.font.size = Pt(11)

    if spec.get("title"):
        h = doc.add_heading(spec["title"], level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta = spec.get("meta", {})
    if meta:
        line = "  ·  ".join(f"{k}: {v}" for k, v in meta.items())
        p = doc.add_paragraph(); r = p.add_run(line)
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    heading_count = 0
    image_paths = []
    for b in spec["blocks"]:
        bt = b["type"]
        if bt == "heading":
            doc.add_heading(b["text"], level=min(max(int(b.get("level", 1)), 1), 4))
            heading_count += 1
        elif bt == "paragraph":
            doc.add_paragraph(b["text"])
        elif bt == "bullets":
            for it in b["items"]:
                doc.add_paragraph(str(it), style="List Bullet")
        elif bt == "quote":
            p = doc.add_paragraph(); r = p.add_run(f"\u201c{b['text']}\u201d")
            r.italic = True
            try:
                p.style = doc.styles["Quote"]
            except KeyError:
                pass
        elif bt == "image":
            if not os.path.exists(b["path"]):
                raise FileNotFoundError(f"image path tidak ada (anti-halu): {b['path']}")
            image_paths.append(b["path"])
            doc.add_picture(b["path"])
            if b.get("caption"):
                cp = doc.add_paragraph(); cr = cp.add_run(b["caption"])
                cr.italic = True; cr.font.size = Pt(9)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif bt == "pagebreak":
            doc.add_page_break()

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)

    text = extract_text(spec)
    return {
        "artifact": out_path,
        "word_count": len(text.split()),
        "heading_count": heading_count,
        "image_paths": image_paths,
        "file_size": os.path.getsize(out_path),
    }
